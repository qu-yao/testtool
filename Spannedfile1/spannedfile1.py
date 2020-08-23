# coding=utf-8
import os
import re
# from time import strftime

import win32com.client
from docx import Document


class Docx:
    def __init__(self, caseid, casedescription, casenature, module):
        self.document = Document()
        self.document.add_heading('案例编号：', 3)
        self.document.add_paragraph(caseid)
        self.document.add_heading('执行人：', 3)
        self.document.add_paragraph("")
        self.document.add_heading('执行时间：', 3)
        self.document.add_paragraph("")
        self.document.add_heading('所属模块：', 3)
        self.document.add_paragraph(module)
        self.document.add_heading('案例描述：', 3)
        self.document.add_paragraph(casedescription)
        self.document.add_heading('案例性质：', 3)
        self.document.add_paragraph(casenature)
        self.document.add_heading('测试步骤：', 3)

    def step(self, step1):
        if len(step1) == 0:
            return
        if type(step1[0]) == list:
            for xx, yy in zip(step1[0], step1[1]):
                if xx == '': continue
                stepstre = '步骤%s' % xx[0], '：', xx, '——>预期结果：', yy
                self.document.add_paragraph(stepstre)
                self.document.add_paragraph("")
        else:
            for str5 in step1:
                self.document.add_paragraph(str5)
                self.document.add_paragraph("")

    def stepone(self, step1, step2, stepname):
        if type(stepname) == float:
            stepname = str(int(stepname))
            if step2 == "":
                str1 = stepname + ":" + step1
                self.document.add_paragraph(str1)
                self.document.add_paragraph("")
            else:
                str1 = stepname + ":" + step1 + "     预期结果     " + step2
                self.document.add_paragraph(str1)
                self.document.add_paragraph("")
        else:
            if step2 == "":
                str1 = step1
                self.document.add_paragraph(str1)
                self.document.add_paragraph("")
            else:
                str1 = step1 + "     预期结果     " + step2
                self.document.add_paragraph(str1)
                self.document.add_paragraph("")

    def save(self, casename, caseid, filepath):

        try:
            self.document.save(filepath + '【' + str(caseid) + '】' + casename[-70:] + '.docx')
        except:
            self.document.save(filepath + '【' + str(caseid) + '】' + casename + '.docx')


def getTest(tdc, caseid):
    testF = tdc.TestFactory
    testFilter = testF.Filter
    testFilter.SetFilter("TS_TEST_ID", caseid)
    testlst = testF.NewList(testFilter.Text)
    # print(testlst.Count)
    # test1 = testlst.Item(1)
    # print(test1.ID)
    # print(test1.ExecStatus)
    for test in testlst:
        return test
    return ""


#   获取步骤list
def getSteplst(theTest):
    dsFact = theTest.DesignStepFactory
    steplst = dsFact.NewList("")
    return steplst


#   获取步骤数据
def getStepData(steplist):
    list1 = []
    if steplist.Count == 0:
        return list1
    for dsstep in steplist:
        if '<html><body>' in dsstep.StepDescription:
            str04 = dsstep.StepDescription
            str004 = str04.replace('<html><body>', '')
            str0004 = str004.replace('</body></html>', '')
            str05 = dsstep.StepExpectedResult
            str005 = str05.replace('<html><body>', '')
            str0005 = str005.replace('</body></html>', '')
            if '<br/>' in str0004:
                description = str0004.split('<br/>')
            else:
                description = str0004.split('<br>')

            if '<br/>' in str0005:
                expectedresult = str0005.split('<br/>')
            else:
                expectedresult = str0005.split('<br>')
            list1.append(description)
            list1.append(expectedresult)
            return list1
        elif '<html>' in dsstep.StepExpectedResult:
            Result_str = dsstep.StepExpectedResult
            Result_str2 = re.search(r'8pt">(.*?)</span>', Result_str).group(1)
            str03 = dsstep.StepName, '：', dsstep.StepDescription, '——>预期结果：', Result_str2
            list1.append(str03)
        else:
            str03 = dsstep.StepName, '：', dsstep.StepDescription, '——>预期结果：', dsstep.StepExpectedResult
            list1.append(str03)

    return list1


def FindTestSets(tdc, path):
    tsTreeMgr = tdc.TestSetTreeManager
    tsFolder = tsTreeMgr.NodeByPath(path)
    tsList = tsFolder.FindTestSets('')
    namelist = []
    idlist = []
    for atestset in tsList:
        print(atestset.name, atestset.id)
        namelist.append(atestset.name)
        idlist.append(atestset.id)
    return idlist


def mainB(filename, url, username, password, domain, project, testSetstr, status, casepath):
    def filepath(str1, filenamepath1):
        x = 1
        if os.path.exists(filenamepath1 + "/%s/" % str1):
            while os.path.exists(filenamepath1 + "/%s/" % (str1 + '-' + str(x))):
                x += 1
            os.mkdir(filenamepath1 + "/%s/" % (str1 + '-' + str(x)))
            return filenamepath1 + "/%s/" % (str1 + '-' + str(x))
        else:
            os.mkdir(filenamepath1 + "/%s/" % str1)
            return filenamepath1 + "/%s/" % str1

    def someset(testSetID):
        TestSetFilter = td.TestSetFactory.Filter  # 获取过滤器
        TestSetFilter["CY_CYCLE_ID"] = testSetID  # 将test set id作为过滤条件
        stList = TestSetFilter.NewList()  # 获取到过滤得到的结果列表，列表中是对象TestSet
        stObject = stList.Item(1)  # 拿到TestSet对象
        TSTestFilter = stObject.TSTestFactory.Filter
        TestSetTestsList = TSTestFilter.NewList()
        print("连接就绪!")
        times = 0
        for test in TestSetTestsList:
            if test.Status != status:
                continue
            if test.HasAttachment:
                continue
            casename = test.Field("TS_NAME")
            casename = casename.replace('\n', '')
            casename = casename.replace('\r', '')
            caseid = test.TestId
            print(casename)
            # playuser = username
            # playtime = strftime('%Y年%m月%d日 %H时%M分')
            module = test.Field("TS_PATH")
            TS_DESCRIPTION = test.Field("TS_DESCRIPTION")
            str01 = TS_DESCRIPTION.replace('<html><body>', '')
            str03 = str01.replace('</body></html>', '')
            str02 = str03.replace('\n', '')
            if '<br' in str02:
                strlist = re.search(r'验证(.*?)<br', str02).group()
                strlist = strlist.replace('<br', '')
            else:
                strlist = str02
            casedescription = strlist

            stepstr = getStepData(getSteplst(getTest(td, caseid)))
            print(stepstr)

            casenature = test.Field("TS_USER_01")

            word = Docx(caseid, casedescription, casenature, module)

            word.step(stepstr)
            word.save(casename, caseid, path_)
            # 检查是否有附件
            # if test.HasAttachment:
            #     print(casename,'有附件')

            times += 1

        return times

    td = win32com.client.Dispatch("TDApiOle80.TDConnection")
    td.InitConnection(url)
    td.Login(username, password)
    td.Connect(domain, project)
    times1 = 0
    if testSetstr == '':
        setlist = FindTestSets(td, casepath)
        for testSetID in setlist:
            path_ = filepath(str(testSetID), filename)
            times1 += someset(testSetID)
    elif '+' in testSetstr:
        setlist = testSetstr.split('+')
        for testSetID in setlist:
            path_ = filepath(str(testSetID), filename)
            times1 += someset(testSetID)
    else:
        path_ = filepath(str(testSetstr), filename)
        times1 += someset(testSetstr)
    s = td.logout
    print(s)
    print("总共生成文件%d个！" % times1)
    return "总共生成文件%d个！" % times1
