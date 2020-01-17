# coding=utf-8
import os
from docx import Document
from Spannedfile.Exclread import ExclRead


def transition(str1):
    list = ["a", "b", "c", "d", "e", "f", "g", "h", "i", "j", "k", "l", "m", "n", "o", "p", "q", "r",
            "s", "t", "u", "v", "w", "x", "y", "z"]
    list1 = ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M", "N", "O", "P", "Q", "R",
             "S", "T", "U", "V", "W", "X", "Y", "Z"]
    if str1 in list:
        data = list.index(str1)
        return data
    elif str1 in list1:
        data = list1.index(str1)
        return data
    else:
        return str1


class Docx():
    def __init__(self, caseid, needid, casedescription, casenature, module):
        self.document = Document()
        self.document.add_heading('案例编号：', 3)
        self.document.add_paragraph(caseid)
        self.document.add_heading('需求编号：', 3)
        self.document.add_paragraph(needid)
        self.document.add_heading('所属模块：', 3)
        self.document.add_paragraph(module)
        self.document.add_heading('案例描述：', 3)
        self.document.add_paragraph(casedescription)
        self.document.add_heading('案例性质：', 3)
        self.document.add_paragraph(casenature)
        self.document.add_heading('测试步骤：', 3)

    def step(self, step1, step2):
        for stepid in step2:
            str3 = ""
            for str1 in stepid:
                if str1 != ".":
                    str3 += str1
                else:
                    break
            indexid = int(str3) - 1
            step1[indexid] = step1[indexid] + "         预期结果         " + stepid
        for str5 in step1:
            self.document.add_paragraph(str5)
            self.document.add_paragraph("")

    def stepone(self, step1, step2, stepname):
        if type(stepname) == float:
            stepname = str(int(stepname))
            if step2 == "":
                str1 = stepname + ":" + step1
                self.document.add_paragraph(str1)
                self.document.add_paragraph("")
            else:
                str1 = stepname + ":" + step1 + "     预期结果     " + step2
                self.document.add_paragraph(str1)
                self.document.add_paragraph("")
        else:
            if step2 == "":
                str1 = step1
                self.document.add_paragraph(str1)
                self.document.add_paragraph("")
            else:
                str1 = step1 + "     预期结果     " + step2
                self.document.add_paragraph(str1)
                self.document.add_paragraph("")

    def save(self, caseid, str1, filenamepath1):
        try:
            self.document.save(filenamepath1 + "/%s/" % str1 + str(caseid) + '.docx')
        except:
            os.mkdir(filenamepath1 + "/%s/" % str1)
            self.document.save(filenamepath1 + "/%s/" % str1 + str(caseid) + '.docx')


def mainA(filename, caseid1, needid1, casedescription1, casenature1, stepid1, resultid1, module1, stepname1,
          filenamepath):
    print(filename, caseid1, needid1, casedescription1, casenature1, stepid1, resultid1, module1, stepname1,
          filenamepath)
    filename = filename
    caseid1 = transition(caseid1)
    needid1 = transition(needid1)
    casedescription1 = transition(casedescription1)
    casenature1 = transition(casenature1)
    stepid1 = transition(stepid1)
    resultid1 = transition(resultid1)
    module1 = transition(module1)
    stepname1 = transition(stepname1)
    filenamepath1 = filenamepath

    exclread = ExclRead(filename, 0)
    sheetnamelist = exclread.sheetname()
    sheetlist = exclread.Excl_sheet()
    for sheetid in sheetlist:
        exclread = ExclRead(filename, sheetid)
        ynum = exclread.Excl_Read_Ynum()
        hb = exclread.merged()
        list11 = exclread.Excl_Read_colx(0)
        if hb == [] and "" != list11[2]:
            print("开始无合并用例生成")
            y = 1
            while y < ynum:
                if exclread.Excl_Read_rowx(y)[stepid1] == "":
                    y += 1
                    continue

                needid = exclread.Excl_Read(y, needid1)
                caseid = exclread.Excl_Read(y, caseid1)
                casedescription = exclread.Excl_Read(y, casedescription1)
                casenature = exclread.Excl_Read(y, casenature1)
                module = exclread.Excl_Read(y, module1)

                if type(needid1) != int:
                    needid = needid1
                if type(caseid1) != int:
                    caseid = caseid1
                if type(casedescription1) != int:
                    casedescription = casedescription1
                if type(casenature1) != int:
                    casenature = casenature1
                if type(module1) != int:
                    module = module1
                d = Docx(caseid, needid, casedescription, casenature, module)
                step1 = exclread.Excl_Read(y, stepid1)
                step2 = exclread.Excl_Read(y, resultid1)
                str1 = step1.splitlines()
                str2 = step2.splitlines()
                d.step(str1, str2)
                sheetname1 = sheetnamelist[sheetid]
                d.save(caseid, sheetname1, filenamepath1)
                y += 1
        else:
            print("开始有合并用例生成")
            y = 1
            while y < ynum:
                if exclread.Excl_Read_rowx(y)[stepid1] == "":
                    y += 1
                    continue
                if exclread.Excl_Read(y, 0) != "":
                    needid = exclread.Excl_Read(y, needid1)
                    caseid = exclread.Excl_Read(y, caseid1)
                    casedescription = exclread.Excl_Read(y, casedescription1)
                    casenature = exclread.Excl_Read(y, casenature1)
                    module = exclread.Excl_Read(y, module1)
                    if type(needid1) != int:
                        needid = needid1
                    if type(caseid1) != int:
                        caseid = caseid1
                    if type(casedescription1) != int:
                        casedescription = casedescription1
                    if type(casenature1) != int:
                        casenature = casenature1
                    if type(module1) != int:
                        module = module1
                    d = Docx(caseid, needid, casedescription, casenature, module)
                    step1 = exclread.Excl_Read(y, stepid1)
                    step2 = exclread.Excl_Read(y, resultid1)
                    stepname = exclread.Excl_Read(y, stepname1)
                    step1 = step1.replace("\n", "")
                    step2 = step2.replace("\n", "")
                    d.stepone(step1, step2, stepname)
                else:
                    step1 = exclread.Excl_Read(y, stepid1)
                    step2 = exclread.Excl_Read(y, resultid1)
                    stepname = exclread.Excl_Read(y, stepname1)
                    step1 = step1.replace("\n", "")
                    step2 = step2.replace("\n", "")
                    d.stepone(step1, step2, stepname)
                y += 1
                sheetname1 = sheetnamelist[sheetid]
                d.save(caseid, sheetname1, filenamepath1)

    print("执行结束")

# mainA()
